from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parent


def set_run_font(run, name: str, size: float, color: str, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    set_run_font(title.add_run("Waypoint DOCX Import Fixture"), "Calibri", 22, "0B2545", True)

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)
    set_run_font(heading.add_run("Acceptance marker"), "Calibri", 16, "2E74B5", True)

    body = document.add_paragraph()
    body.add_run("WAYPOINT_DOCX_42B7 is a disposable Windows acceptance marker. ")
    body.add_run("It verifies Word import, local extraction, indexing, search, provenance, reindex, and deletion behavior.")

    source = document.add_paragraph()
    source.add_run("Created only for QA on 2026-08-06. No production or personal data is present.")

    document.save(ROOT / "QA-Waypoint-DOCX-42B7.docx")


def build_pdf() -> None:
    output = ROOT / "QA-Waypoint-PDF-73C1.pdf"
    pdf = canvas.Canvas(str(output), pagesize=letter)
    width, height = letter
    pdf.setTitle("Waypoint PDF Import Fixture")
    pdf.setFillColor(HexColor("#0B2545"))
    pdf.setFont("Helvetica-Bold", 22)
    pdf.drawString(72, height - 90, "Waypoint PDF Import Fixture")
    pdf.setFillColor(HexColor("#2E74B5"))
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(72, height - 130, "Acceptance marker")
    pdf.setFillColor(HexColor("#1F2937"))
    pdf.setFont("Helvetica", 11)
    pdf.drawString(72, height - 160, "WAYPOINT_PDF_73C1 is a disposable Windows acceptance marker.")
    pdf.drawString(72, height - 180, "It verifies PDF import, local extraction, indexing, search, provenance, reindex, and deletion.")
    pdf.drawString(72, height - 215, "Created only for QA on 2026-08-06. No production or personal data is present.")
    pdf.setStrokeColor(HexColor("#D7DEE8"))
    pdf.line(72, height - 235, width - 72, height - 235)
    pdf.save()


if __name__ == "__main__":
    build_docx()
    build_pdf()
